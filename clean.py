import jieba.posseg as pseg
import jieba
from docx import Document

"""
可以把违禁词替换词"xxx"和补齐的空白符<pad>添加到word2id的字典中
"""

word2id = {}
id2word = {}
tag2id = {}
id2tag = {}
word2tag = {}
# 读取停用词
stop_words = []
for line in open("stop_words.txt",encoding="UTF-8"):
    stop_words.append(line.strip())

# 读取违禁词
prohibited_words = []
for line in open("prohibited_words.txt",encoding="UTF-8"):
    prohibited_words.append(line.strip())

# print(stop_words[:90],prohibited_words)

# 读取文章
path = "pocess_data.docx"
document = Document(path)
#词列表
word_list = []
tag_list = []
for paragraph in document.paragraphs:
    text = paragraph.text
    screen_text = ""
    for i in text:
        if i >= u'\u4e00' and i <= u'\u9fa5':
            screen_text += i
    words = pseg.cut(screen_text)
    for w in words:
        if w.word in stop_words:
            continue
        elif w.word in prohibited_words:
            continue
        else:
            word_list.append(w.word)
            tag_list.append(w.flag)
            word2tag[w.word] = w.flag


# 词频 dict
word_frequency = {}
for w in word_list:
    if not w in word_frequency:
        word_frequency[w] = 1
    else:
        word_frequency[w] += 1


# 配置项
# 句子长度
setence_length = int(input("请输入句子长度："))
#setence_length = 5
# 筛选的词频
screen_num = int(input("请输入筛选的词频："))


# 筛掉低频词
for w in word_list:
    try:
        if word_frequency[w] and word_frequency[w]<=screen_num:
            word_frequency.pop(w)
    except:
        continue
# print(word_frequency)

# 将word 和 tag 都转化为id
for w in word_frequency:
    if not w in word2id:
        word2id[w] = len(word2id)
        id2word[len(id2word)] = w
# 添加空词<pad>
word2id["<pad>"] = len(word2id)
id2word[len(id2word)] = "<pad>"
# 添加违禁词替换词
word2id["XXX"] = len(word2id)
id2word[len(id2word)] = "XXX"


for t in tag_list:
    if not t in tag2id:
        tag2id[t] = len(tag2id)
        id2tag[len(id2tag)] = t
# 添加空词<pad>的tag，也为<pad>
tag2id["<pad>"] = len(tag2id)
id2tag[len(id2tag)] = "<pad>"
# 添加违禁词tag
tag2id["XXX"] = len(tag2id)
id2tag[len(id2tag)] = "XXX"

print(word2id)
#pad补齐的函数
def pad(s_list,setence_length):
    s_length = len(s_list)
    if(s_length>=setence_length):
        return s_list[:setence_length-1]
    else:
        while(s_length < setence_length):
            #添加<pad>的id 和 tag_id
            s_list.append("("+str(word2id["<pad>"])+","+str(tag2id["<pad>"])+")")
            s_length +=1
        return s_list


# 输出结果
result = []
for paragraph in document.paragraphs:
    #该行替换列表
    sentence_to_list = []

    text = paragraph.text
    screen_text = ""
    for i in text:
        if i >= u'\u4e00' and i <= u'\u9fa5':
            screen_text += i
    words = pseg.cut(screen_text)
    for w in words:
        # 先声明词id 和 词性id
        w_id = ""
        tag_id = ""
        # 如果该词为停用词
        if w.word in stop_words:
            continue
        # 如果该词为违禁词
        elif w.word in prohibited_words:
            w_id = str(word2id["XXX"])
            tag_id = str(tag2id["XXX"])
            sentence_to_list.append("("+w_id+","+tag_id+")")
            print("有查询到违禁词")
            continue
        elif w.word in word2id:
            w_id = str(word2id[w.word])
            tag_id = str(tag2id[word2tag[w.word]])
            sentence_to_list.append("(" + w_id + "," + tag_id + ")")
            continue
        else:
            print(w.word)
    #进行长度控制
    pad(sentence_to_list,setence_length)
    # 添加每个sentence的list到结果的list里面
    result.append(sentence_to_list)
print(result)