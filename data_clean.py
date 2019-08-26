import jieba.posseg as pseg
from docx import Document
from collections import Counter

"""
负责清洗数据的类，根据配置删去低频词，停用词，标点符号，空行等等只保留中文文本信息并转化
成对应的id号和词标签id号，固定句长，短补长截。
"""
class Clean:
    """
    初始化，对应参数分别为：固定句子长度，删去词的词频，读取word文档的路径
    """
    def __init__(self,setence_length=5,screen_num=1,file_path="pocess_data.docx"):
        self.sentence_length = setence_length
        self.screen_num = screen_num
        self.file_path = file_path
        self.data2id()

    """
    获取各个数据对应的id字典
    """
    def data2id(self):
        # 中文词word对应id
        word2id = {}
        id2word = {}
        # 中文词的词性tag对应的id
        tag2id = {}
        id2tag = {}
        # 中文词word和词性tag对应的字典
        word2tag = {}

        # 读取停用词
        stop_words = []
        for line in open("stop_words.txt", encoding="UTF-8"):
            stop_words.append(line.strip())
        stop_words = set(stop_words)

        # 读取违禁词
        prohibited_words = []
        for line in open("prohibited_words.txt", encoding="UTF-8"):
            prohibited_words.append(line.strip())
        prohibited_words = set(prohibited_words)
        # print(stop_words[:90],prohibited_words)

        # 读取文章
        document = Document(self.file_path)
        # 词列表
        word_list = []
        tag_list = []
        for paragraph in document.paragraphs:
            text = paragraph.text
            screen_text = ""
            for i in text:
                if i >= u'\u4e00' and i <= u'\u9fa5':
                    screen_text += i
            words = pseg.cut(screen_text)
            for w in words:
                if w.word in stop_words:
                    continue
                elif w.word in prohibited_words:
                    continue
                else:
                    word_list.append(w.word)
                    tag_list.append(w.flag)
                    word2tag[w.word] = w.flag

        # 词频 dict
        # word_frequency = {}
        # for w in word_list:
        #     if w not in word_frequency:
        #         word_frequency[w] = 1
        #     else:
        #         word_frequency[w] += 1
        word_frequency = Counter(word_list)
        # 筛掉低频词
        # word2id = {items[0]:idx for idx,items in enumerate(word_frequency) if items[1]>self.screen_num }
        # for w in word_list:
        #     try:
        #         if word_frequency[w] and word_frequency[w] <= screen_num:
        #             word_frequency.pop(w)
        #     except:
        #         continue

        # 筛掉低频词
        vocab_list = [word for word in word_frequency if word_frequency[word] > self.screen_num]
        # 将word 和 tag 都转化为id
        for w in vocab_list:
            if w not in word2id:
                word2id[w] = len(word2id)
                id2word[len(id2word)] = w
        # 添加空词<pad>
        word2id["<pad>"] = len(word2id)
        id2word[len(id2word)] = "<pad>"
        # 添加违禁词替换词
        word2id["XXX"] = len(word2id)
        id2word[len(id2word)] = "XXX"
        for t in tag_list:
            if t not in tag2id:
                tag2id[t] = len(tag2id)
                id2tag[len(id2tag)] = t
        # 添加空词<pad>的tag，也为<pad>
        tag2id["<pad>"] = len(tag2id)
        id2tag[len(id2tag)] = "<pad>"
        # 添加违禁词tag
        tag2id["XXX"] = len(tag2id)
        id2tag[len(id2tag)] = "XXX"
        # print(word2id)
        self.stop_words = set(stop_words)
        self.prohibited_words = set(prohibited_words)
        self.word2id = word2id
        self.id2word = id2word
        self.tag2id = tag2id
        self.id2tag = id2tag
        self.word2tag = word2tag
        return


    """
    固定句长函数，如果这个句子长度大于规定句长则截断，如不足就补齐<pad>
    """
    def pad(self,s_list,setence_length):
        s_length = len(s_list)
        if(s_length>=setence_length):
            return s_list[:setence_length-1]
        else:
            while(s_length < setence_length):
                #添加<pad>的id 和 tag_id
                s_list.append("("+str(self.word2id["<pad>"])+","+str(self.tag2id["<pad>"])+")")
                s_length +=1
            return s_list

    """
    获取最终结果，为一个list，其中每一项都是固定句长的一行的list，内部分别是word_id和tag_id
    """
    def get_result(self):
        result = []
        document = Document(self.file_path)
        for paragraph in document.paragraphs:
            # 该行替换列表
            sentence_to_list = []
            text = paragraph.text
            # 去除空行
            if(text.strip() == ""):
                continue
            screen_text = ""
            for i in text:
                if i >= u'\u4e00' and i <= u'\u9fa5':
                    screen_text += i
            # print(screen_text)
            words = pseg.cut(screen_text)
            for w in words:
                # 声明词id 和 词性id
                w_id = ""
                tag_id = ""
                # 如果该词为停用词
                if w.word in self.stop_words:
                    continue
                # 如果该词为违禁词
                elif w.word in self.prohibited_words:
                    w_id = str(self.word2id["XXX"])
                    tag_id = str(self.tag2id["XXX"])
                    sentence_to_list.append("(" + w_id + "," + tag_id + ")")
                    print("有查询到违禁词")
                    continue
                elif w.word in self.word2id:
                    w_id = str(self.word2id[w.word])
                    tag_id = str(self.tag2id[w.flag])
                    sentence_to_list.append("(" + w_id + "," + tag_id + ")")
                    continue
                else:
                    # 这些都是低频词，不需要处理，直接pass
                    pass
            # 进行长度控制
            self.pad(sentence_to_list, setence_length)
            # 添加每个sentence的list到结果的list里面
            result.append(sentence_to_list)
        return result

    """
    将结果解析回中文词和词性
    """
    def parse(self,result_list):
        parse_list = []
        for sentence_list in result_list:
            parse_sentence = []
            for i in sentence_list:
                temp = i[1:-1].split(",")
                w_id = int(temp[0])
                t_id = int(temp[1])
                word = self.id2word[w_id]
                tag = self.id2tag[t_id]
                item = "("+word+","+tag+")"
                parse_sentence.append(item)
            parse_list.append(parse_sentence)
        return parse_list

# 配置
# 句子长度
setence_length = int(input("请输入句子长度："))
# 筛选的词频
screen_num = int(input("请输入过滤的词频："))

c = Clean(setence_length,screen_num)
result = c.get_result()
print(result)
parse = c.parse(result)
print(parse)
